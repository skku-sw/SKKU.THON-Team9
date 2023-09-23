import boto3
from request_model import (
    aws_access_key_id,
    aws_secret_access_key,
    region_name,
    s3_bucket_name,
    image_path,
)
from flask import (
    Blueprint,
    jsonify,
    request,
    session,
    redirect,
    url_for,
    render_template_string,
)
import os
import subprocess

# apt-get install libreoffice
# sudo apt-get install unoconv
from db_model import UserInfo, to_dict, DoctorInfo, MedicalHistory
from flask_jwt_extended import create_access_token, jwt_required, get_jwt_identity
from werkzeug.security import check_password_hash
from request_model import (
    login_request_model,
    register_request_model,
    upload_request_model,
)
from docx import Document
from docx2pdf import convert

# boto3 클라이언트 생성
s3_client = boto3.client(
    "s3",
    aws_access_key_id=aws_access_key_id,
    aws_secret_access_key=aws_secret_access_key,
    region_name=region_name,
)

pdf_api = Blueprint("pdf_api", __name__)


def docx_to_pdf(input_path, output_path):
    subprocess.check_call(["unoconv", "-f", "pdf", "-o", output_path, input_path])


@pdf_api.route("/upload_diagnosis", methods=["POST"])
def generate_and_upload():
    try:
        data = request.get_json()
        # 1. 입력받은 데이터 유효성 검사
        if not all(field in data for field in upload_request_model):
            return jsonify({"msg": "Missing or invalid data."}), 400

        # 2. DB에 넣기
        user = UserInfo.query.filter_by(patient_id=data["patient_id"]).first()
        if not user:
            return jsonify({"msg": "User Not Found"}), 404

        doctor = DoctorInfo.query.filter_by(
            license_number=data["license_number"]
        ).first()
        if not doctor:
            doctor = DoctorInfo(
                license_number=data["license_number"],
                medical_institution=data["medical_institution"],
                doctor_name=data["doctor_name"],
            )

        medical_history = MedicalHistory(
            patient_id=data["patient_id"],
            license_number=data["license_number"],
            diagnosis_code=data["diagnosis_code"],
            onset_date=data["onset_date"],
            diagnosis_date=data["diagnosis_date"],
            prognosis=data["prognosis"],
            hospitalization_dates=data["hospitalization_dates"],
        )

        # 3. Word template에서 객체 가져오기
        doc = Document("template.docx")

        # 4. 템플릿의 Placeholder를 DB에서 가져온 값으로 대체
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{0}" in paragraph.text:
                            text = paragraph.text.replace("{0}", user.full_name)
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{1}" in paragraph.text:
                            text = paragraph.text.replace("{1}", user.gender)
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{2}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{2}", str(user.date_of_birth)
                            )
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{3}" in paragraph.text:
                            text = paragraph.text.replace("{3}", user.address)
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{4}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{4}", medical_history.diagnosis_code
                            )
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{5}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{5}", str(medical_history.onset_date)
                            )
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{6}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{6}", str(medical_history.diagnosis_date)
                            )
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{7}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{7}", medical_history.prognosis
                            )
                            paragraph.clear()
                            paragraph.add_run(text)
                        elif "{8}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{8}", str(medical_history.diagnosis_date)
                            )
                            paragraph.clear()
                            run = paragraph.add_run(text)
                            run.bold = True
                        elif "{9}" in paragraph.text:
                            text = paragraph.text.replace(
                                "{9}", doctor.medical_institution
                            )
                            paragraph.clear()
                            run = paragraph.add_run(text)
                            run.bold = True
                        elif "{10}" in paragraph.text:
                            text = paragraph.text.replace("{10}", doctor.license_number)
                            paragraph.clear()
                            run = paragraph.add_run(text)
                            run.bold = True
                        elif "{11}" in paragraph.text:
                            text = paragraph.text.replace("{11}", doctor.doctor_name)
                            paragraph.clear()
                            run = paragraph.add_run(text)
                            run.bold = True
                        elif "{12}" in paragraph.text:
                            text = paragraph.text.replace("{12}", user.patient_id)
                            paragraph.clear()
                            paragraph.add_run(text)

        # 5. Word 문서 임시 저장
        temp_docx_path = "temp.docx"
        doc.save(temp_docx_path)
        # 6. Word를 PDF로 변환
        temp_pdf_path = f"{medical_history.patient_id}_{medical_history.license_number}_{medical_history.diagnosis_date}.pdf"
        convert(temp_docx_path, temp_pdf_path)
        # docx_to_pdf(temp_docx_path, temp_pdf_path)
        # 7. PDF를 S3에 업로드
        with open(temp_pdf_path, "rb") as file:
            s3_client.upload_fileobj(file, s3_bucket_name, "image/" + temp_pdf_path)

        # DB 저장
        medical_history.filename = image_path + "image/" + temp_pdf_path
        from app import db_session

        db_session.merge(doctor)
        db_session.commit()
        db_session.merge(medical_history)
        db_session.commit()

        # 8. 파일 기록 지우기
        os.remove(temp_docx_path)
        os.remove(temp_pdf_path)
        return jsonify({"msg": "Upload Success"}), 200
    except Exception as e:
        return jsonify({"msg": str(e)}), 401
